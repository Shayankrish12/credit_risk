"""PDF and DOCX report generation for credit monitoring notes."""
import io
from datetime import datetime, timezone
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak)
from docx import Document
from docx.shared import Pt, RGBColor, Inches


def _risk_color(category: str):
    return {
        "low": colors.HexColor("#38A169"),
        "moderate": colors.HexColor("#002FA7"),
        "high": colors.HexColor("#D69E2E"),
        "critical": colors.HexColor("#E53E3E"),
    }.get(category, colors.HexColor("#0A0A0A"))


def build_pdf(borrower: dict, factors: list, signals: list, notes: list, sales: list, bank: list, repayments: list) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('title', parent=styles['Heading1'], fontSize=20, spaceAfter=12)
    h2 = ParagraphStyle('h2', parent=styles['Heading2'], fontSize=13, spaceBefore=14, spaceAfter=6, textColor=colors.HexColor("#0A0A0A"))
    body = styles['BodyText']

    story = []
    story.append(Paragraph("CREDIT MONITORING NOTE", title_style))
    story.append(Paragraph(f"Generated on {datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M UTC')}", body))
    story.append(Spacer(1, 0.3*cm))

    # Borrower overview
    story.append(Paragraph("BORROWER OVERVIEW", h2))
    overview_data = [
        ["Business Name", borrower.get("business_name", "")],
        ["Sector", borrower.get("sector", "")],
        ["Location", borrower.get("location", "")],
        ["Loan Type", borrower.get("loan_type", "")],
        ["Loan Amount", f"INR {borrower.get('loan_amount', 0):,.0f}"],
        ["Outstanding", f"INR {borrower.get('outstanding_amount', 0):,.0f}"],
        ["Sanction Date", borrower.get("sanction_date", "")],
        ["GST Number", borrower.get("gst_number", "")],
        ["Contact Person", borrower.get("contact_person", "")],
    ]
    t = Table(overview_data, colWidths=[5*cm, 11*cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
        ("FONTSIZE", (0,0), (-1,-1), 10),
        ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#E5E7EB")),
        ("BACKGROUND", (0,0), (0,-1), colors.HexColor("#F9FAFB")),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("PADDING", (0,0), (-1,-1), 6),
    ]))
    story.append(t)

    # Risk Score
    story.append(Paragraph("RISK ASSESSMENT", h2))
    risk_data = [
        ["Risk Score", f"{borrower.get('risk_score', 0):.1f} / 100"],
        ["Risk Category", borrower.get("risk_category", "").upper()],
    ]
    rt = Table(risk_data, colWidths=[5*cm, 11*cm])
    rt.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 11),
        ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#E5E7EB")),
        ("BACKGROUND", (0,0), (0,-1), colors.HexColor("#F9FAFB")),
        ("TEXTCOLOR", (1,1), (1,1), _risk_color(borrower.get("risk_category", "low"))),
        ("PADDING", (0,0), (-1,-1), 6),
    ]))
    story.append(rt)

    # Financial trend summary
    story.append(Paragraph("FINANCIAL TREND SUMMARY", h2))
    if sales:
        recent_sales = sum(s["amount"] for s in sales[-3:]) / max(min(3, len(sales)), 1)
        prior_sales = sum(s["amount"] for s in sales[:-3]) / max(len(sales) - 3, 1) if len(sales) > 3 else recent_sales
        sales_change = ((recent_sales - prior_sales) / prior_sales * 100) if prior_sales > 0 else 0
        story.append(Paragraph(f"Average recent sales: INR {recent_sales:,.0f} per month ({sales_change:+.1f}% vs earlier period)", body))
    if bank:
        recent_bank = sum(b["balance"] for b in bank[-3:]) / max(min(3, len(bank)), 1)
        story.append(Paragraph(f"Average recent bank balance: INR {recent_bank:,.0f}", body))
    if repayments:
        bounced = sum(1 for r in repayments if r.get("status") == "bounced")
        delayed = sum(1 for r in repayments if r.get("status") == "delayed")
        story.append(Paragraph(f"Repayment history: {bounced} bounces, {delayed} delays out of {len(repayments)} cycles", body))

    # Contributing factors
    story.append(Paragraph("TOP CONTRIBUTING RISK FACTORS", h2))
    if factors:
        factor_rows = [["Factor", "Impact", "Detail"]]
        for f in factors[:8]:
            factor_rows.append([f.get("factor", ""), f"{f.get('impact', 0)}", f.get("detail", "")])
        ft = Table(factor_rows, colWidths=[5*cm, 2*cm, 9*cm])
        ft.setStyle(TableStyle([
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#F3F4F6")),
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#E5E7EB")),
            ("PADDING", (0,0), (-1,-1), 5),
        ]))
        story.append(ft)
    else:
        story.append(Paragraph("No major risk factors identified.", body))

    # Warning signals
    story.append(Paragraph("KEY WARNING SIGNALS", h2))
    if signals:
        sig_rows = [["Signal", "Severity", "Explanation"]]
        for s in signals[:10]:
            sig_rows.append([s.get("signal_type", "").replace("_", " ").title(),
                             s.get("severity", "").upper(), s.get("explanation", "")])
        st = Table(sig_rows, colWidths=[4*cm, 2.5*cm, 9.5*cm])
        st.setStyle(TableStyle([
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#F3F4F6")),
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("GRID", (0,0), (-1,-1), 0.25, colors.HexColor("#E5E7EB")),
            ("PADDING", (0,0), (-1,-1), 5),
        ]))
        story.append(st)
    else:
        story.append(Paragraph("No active warning signals.", body))

    # Analyst notes
    story.append(Paragraph("ANALYST OBSERVATIONS", h2))
    if notes:
        for n in notes[:5]:
            story.append(Paragraph(f"<b>{n.get('created_by_name', 'Analyst')}</b> ({n.get('created_at','')[:10]}): {n.get('content','')}", body))
            story.append(Spacer(1, 0.15*cm))
    else:
        story.append(Paragraph("No analyst observations recorded.", body))

    # Recommended action
    story.append(Paragraph("RECOMMENDED ACTION", h2))
    cat = borrower.get("risk_category", "low")
    rec = {
        "low": "Continue standard monitoring. Quarterly review sufficient.",
        "moderate": "Move to monthly monitoring. Request updated financial statements.",
        "high": "Schedule borrower meeting. Reassess collateral and require monthly cash flow updates.",
        "critical": "Escalate to credit committee. Initiate recovery / restructuring discussion immediately.",
    }.get(cat, "Continue monitoring.")
    story.append(Paragraph(rec, body))

    # Follow-up questions
    story.append(Paragraph("FOLLOW-UP QUESTIONS", h2))
    questions = [
        "Can borrower share latest 3 months bank statements and GST returns?",
        "What is the reason for any recent sales decline or volatility?",
        "Are there receivable concentration risks with top customers?",
        "Does the borrower foresee additional working capital needs?",
        "Is there any pending litigation or adverse news to disclose?",
    ]
    for q in questions:
        story.append(Paragraph(f"&bull; {q}", body))

    # Final monitoring status
    story.append(Paragraph("FINAL MONITORING STATUS", h2))
    status_map = {"low": "STANDARD", "moderate": "ENHANCED MONITORING", "high": "SPECIAL MENTION", "critical": "WATCH LIST"}
    story.append(Paragraph(f"<b>{status_map.get(cat, 'STANDARD')}</b>", body))

    doc.build(story)
    return buffer.getvalue()


def build_docx(borrower: dict, factors: list, signals: list, notes: list, sales: list, bank: list, repayments: list) -> bytes:
    doc = Document()
    title = doc.add_heading("Credit Monitoring Note", 0)
    doc.add_paragraph(f"Generated on {datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M UTC')}")

    doc.add_heading("Borrower Overview", 1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in [
        ("Business Name", borrower.get("business_name", "")),
        ("Sector", borrower.get("sector", "")),
        ("Location", borrower.get("location", "")),
        ("Loan Type", borrower.get("loan_type", "")),
        ("Loan Amount", f"INR {borrower.get('loan_amount', 0):,.0f}"),
        ("Outstanding", f"INR {borrower.get('outstanding_amount', 0):,.0f}"),
        ("Sanction Date", borrower.get("sanction_date", "")),
        ("GST Number", borrower.get("gst_number", "")),
        ("Contact Person", borrower.get("contact_person", "")),
    ]:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = str(v)

    doc.add_heading("Risk Assessment", 1)
    p = doc.add_paragraph()
    p.add_run(f"Risk Score: {borrower.get('risk_score', 0):.1f}/100  |  Category: ").bold = True
    cat_run = p.add_run(borrower.get("risk_category", "").upper())
    cat_run.bold = True

    doc.add_heading("Financial Trend Summary", 1)
    if sales:
        recent = sum(s["amount"] for s in sales[-3:]) / max(min(3, len(sales)), 1)
        doc.add_paragraph(f"Average recent sales: INR {recent:,.0f}/month")
    if bank:
        recent_b = sum(b["balance"] for b in bank[-3:]) / max(min(3, len(bank)), 1)
        doc.add_paragraph(f"Average recent bank balance: INR {recent_b:,.0f}")
    if repayments:
        bounced = sum(1 for r in repayments if r.get("status") == "bounced")
        delayed = sum(1 for r in repayments if r.get("status") == "delayed")
        doc.add_paragraph(f"Repayment history: {bounced} bounces, {delayed} delays out of {len(repayments)} cycles")

    doc.add_heading("Top Contributing Risk Factors", 1)
    for f in factors[:8]:
        doc.add_paragraph(f"{f.get('factor','')} (impact +{f.get('impact',0)}): {f.get('detail','')}", style="List Bullet")

    doc.add_heading("Key Warning Signals", 1)
    if signals:
        for s in signals[:10]:
            doc.add_paragraph(f"[{s.get('severity','').upper()}] {s.get('signal_type','').replace('_',' ').title()}: {s.get('explanation','')}", style="List Bullet")
    else:
        doc.add_paragraph("No active warning signals.")

    doc.add_heading("Analyst Observations", 1)
    if notes:
        for n in notes[:5]:
            doc.add_paragraph(f"{n.get('created_by_name','Analyst')} ({n.get('created_at','')[:10]}): {n.get('content','')}", style="List Bullet")
    else:
        doc.add_paragraph("No analyst observations recorded.")

    doc.add_heading("Recommended Action", 1)
    cat = borrower.get("risk_category", "low")
    rec = {
        "low": "Continue standard monitoring. Quarterly review sufficient.",
        "moderate": "Move to monthly monitoring. Request updated financial statements.",
        "high": "Schedule borrower meeting. Reassess collateral and require monthly cash flow updates.",
        "critical": "Escalate to credit committee. Initiate recovery / restructuring discussion immediately.",
    }.get(cat, "Continue monitoring.")
    doc.add_paragraph(rec)

    doc.add_heading("Follow-up Questions", 1)
    for q in [
        "Can borrower share latest 3 months bank statements and GST returns?",
        "What is the reason for any recent sales decline or volatility?",
        "Are there receivable concentration risks with top customers?",
        "Does the borrower foresee additional working capital needs?",
        "Is there any pending litigation or adverse news to disclose?",
    ]:
        doc.add_paragraph(q, style="List Bullet")

    doc.add_heading("Final Monitoring Status", 1)
    status_map = {"low": "STANDARD", "moderate": "ENHANCED MONITORING", "high": "SPECIAL MENTION", "critical": "WATCH LIST"}
    doc.add_paragraph(status_map.get(cat, "STANDARD")).runs[0].bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
